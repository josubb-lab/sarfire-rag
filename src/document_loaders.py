"""
MultiFormatLoader: carga unificada (PDF, DOCX, TXT, HTML, CSV) para SARFIRE-RAG.

Convierte cada fuente en LangChain Document; para indexar con DocumentChunker existente,
usar langchain_documents_to_pdfloader_documents().
"""
from __future__ import annotations

import logging
from collections import Counter
from pathlib import Path
from typing import Dict, List, Sequence

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.documents import Document

logger = logging.getLogger(__name__)

_PYPDF_LOADER = None
_TEXT_LOADER = None


def _get_pdf_loader_class():
    global _PYPDF_LOADER
    if _PYPDF_LOADER is not None:
        return _PYPDF_LOADER
    try:
        from langchain_community.document_loaders import PyPDFLoader as _C

        _PYPDF_LOADER = _C
    except ImportError:
        from langchain.document_loaders import PyPDFLoader as _L

        _PYPDF_LOADER = _L
    return _PYPDF_LOADER


def _get_text_loader_class():
    global _TEXT_LOADER
    if _TEXT_LOADER is not None:
        return _TEXT_LOADER
    try:
        from langchain_community.document_loaders import TextLoader as _C

        _TEXT_LOADER = _C
    except ImportError:
        from langchain.document_loaders import TextLoader as _L

        _TEXT_LOADER = _L
    return _TEXT_LOADER


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".html", ".htm", ".csv"}


class MultiFormatLoader:
    """Carga documentos en varios formatos; cada método devuelve List[Document]."""

    def __init__(self) -> None:
        self._splitter = RecursiveCharacterTextSplitter(
            chunk_size=500,
            chunk_overlap=100,
            separators=["\n\n", "\n", ".", " "],
        )

    def _split_and_index(self, docs: List[Document]) -> List[Document]:
        if not docs:
            return []
        chunks = self._splitter.split_documents(docs)
        for i, d in enumerate(chunks):
            d.metadata["chunk_index"] = i
        return chunks

    def load_pdf(self, path: str | Path) -> List[Document]:
        path = Path(path).resolve()
        fmt = path.suffix.lower()
        PyPDFLoader = _get_pdf_loader_class()
        loader = PyPDFLoader(str(path))
        docs = loader.load()
        for d in docs:
            d.metadata["source"] = str(path)
            d.metadata["format"] = fmt
            if "page" in d.metadata:
                d.metadata["chunk_index"] = int(d.metadata["page"])
        return docs

    def load_docx(self, path: str | Path) -> List[Document]:
        path = Path(path).resolve()
        fmt = path.suffix.lower()
        try:
            from docx import Document as DocxDocument
        except ImportError as e:
            logger.exception("[LOADER] python-docx no disponible: %s", e)
            return []

        try:
            ddoc = DocxDocument(str(path))
            parts = [p.text.strip() for p in ddoc.paragraphs if p.text and p.text.strip()]
            body = "\n\n".join(parts)
            meta = {"source": str(path), "format": fmt, "chunk_index": 0}
            return self._split_and_index([Document(page_content=body, metadata=meta)])
        except Exception as e:
            logger.exception("[LOADER] Error leyendo DOCX %s: %s", path.name, e)
            return []

    def load_txt(self, path: str | Path) -> List[Document]:
        path = Path(path).resolve()
        fmt = path.suffix.lower()
        TextLoader = _get_text_loader_class()
        try:
            loader = TextLoader(str(path), encoding="utf-8")
            docs = loader.load()
            for d in docs:
                d.metadata["source"] = str(path)
                d.metadata["format"] = fmt
            return self._split_and_index(docs)
        except UnicodeDecodeError:
            try:
                loader = TextLoader(str(path), encoding="latin-1")
                docs = loader.load()
                for d in docs:
                    d.metadata["source"] = str(path)
                    d.metadata["format"] = fmt
                return self._split_and_index(docs)
            except Exception as e:
                logger.exception("[LOADER] Error leyendo TXT %s: %s", path.name, e)
                return []
        except Exception as e:
            logger.exception("[LOADER] Error leyendo TXT %s: %s", path.name, e)
            return []

    def load_html(self, path: str | Path) -> List[Document]:
        path = Path(path).resolve()
        fmt = path.suffix.lower()
        try:
            from bs4 import BeautifulSoup
        except ImportError as e:
            logger.exception("[LOADER] beautifulsoup4 no disponible: %s", e)
            return []

        try:
            raw = path.read_text(encoding="utf-8", errors="replace")
            soup = BeautifulSoup(raw, "html.parser")
            for tag in soup(["script", "style", "noscript"]):
                tag.decompose()
            text = soup.get_text(separator="\n", strip=True)
            meta = {"source": str(path), "format": fmt, "chunk_index": 0}
            return self._split_and_index([Document(page_content=text, metadata=meta)])
        except Exception as e:
            logger.exception("[LOADER] Error leyendo HTML %s: %s", path.name, e)
            return []

    def load_csv(self, path: str | Path) -> List[Document]:
        path = Path(path).resolve()
        fmt = path.suffix.lower()
        try:
            import pandas as pd
        except ImportError as e:
            logger.exception("[LOADER] pandas no disponible: %s", e)
            return []

        try:
            df = pd.read_csv(path)
            docs: List[Document] = []
            for idx, row in df.iterrows():
                lines = []
                meta: Dict = {
                    "source": str(path),
                    "format": fmt,
                    "chunk_index": int(idx),
                }
                for col in df.columns:
                    val = row[col]
                    if pd.notna(val):
                        lines.append(f"{col}: {val}")
                        meta[str(col)] = str(val)
                content = "\n".join(lines)
                docs.append(Document(page_content=content, metadata=meta))
            return docs
        except Exception as e:
            logger.exception("[LOADER] Error leyendo CSV %s: %s", path.name, e)
            return []

    def load_all(self, directory: str | Path) -> List[Document]:
        directory = Path(directory).resolve()
        if not directory.is_dir():
            logger.error("[LOADER] No es un directorio: %s", directory)
            return []

        combined: List[Document] = []
        files = sorted(
            f
            for f in directory.iterdir()
            if f.is_file() and not f.name.startswith(".") and f.suffix.lower() in SUPPORTED_EXTENSIONS
        )

        for fp in files:
            ext = fp.suffix.lower()
            try:
                if ext == ".pdf":
                    part = self.load_pdf(fp)
                elif ext == ".docx":
                    part = self.load_docx(fp)
                elif ext == ".txt":
                    part = self.load_txt(fp)
                elif ext in (".html", ".htm"):
                    part = self.load_html(fp)
                elif ext == ".csv":
                    part = self.load_csv(fp)
                else:
                    continue
                n = len(part)
                logger.info("[LOADER] Cargando %s → %s chunks", fp.name, n)
                combined.extend(part)
            except Exception as e:
                logger.exception("[LOADER] Fallo procesando %s: %s", fp.name, e)

        return combined


def langchain_documents_to_pdfloader_documents(documents: Sequence[Document]) -> List[Dict]:
    """
    Agrupa Document de LangChain en la estructura dict que consume DocumentChunker
    (misma forma que `loaders.pdf_loader.PDFLoader`).
    """
    from collections import defaultdict

    groups: dict[str, List[Document]] = defaultdict(list)
    for d in documents:
        src = str(d.metadata.get("source", ""))
        groups[src].append(d)

    legacy: List[Dict] = []
    for source in sorted(groups.keys()):
        items = groups[source]
        path = Path(source)
        fmt = (items[0].metadata.get("format") or path.suffix).lower()
        if fmt == ".pdf":
            items.sort(key=lambda x: int(x.metadata.get("page", x.metadata.get("chunk_index", 0))))
        else:
            items.sort(key=lambda x: int(x.metadata.get("chunk_index", 0)))

        filename = path.name
        pages = []
        for i, doc in enumerate(items):
            if fmt == ".pdf":
                page_num = int(doc.metadata.get("page", i)) + 1
            else:
                page_num = i + 1
            pages.append({"page_num": page_num, "text": doc.page_content})

        nonempty = [p for p in pages if (p.get("text") or "").strip()]
        legacy.append(
            {
                "metadata": {
                    "filename": filename,
                    "num_pages": len(pages),
                    "num_pages_with_text": len(nonempty),
                    "path": source,
                },
                "pages": pages,
            }
        )

    return legacy


def main() -> None:
    logging.basicConfig(level=logging.INFO, format="%(message)s")
    root = Path(__file__).resolve().parent.parent
    raw_dir = root / "data" / "raw"

    loader = MultiFormatLoader()
    docs = loader.load_all(raw_dir)

    by_fmt = Counter((d.metadata.get("format") or "?") for d in docs)
    uniq_formats = sorted(by_fmt.keys())
    print("\n--- Estadísticas ---")
    for fmt in uniq_formats:
        print(f"  {fmt}: {by_fmt[fmt]} documentos (LangChain)")
    print(f"Total: {len(docs)} documentos de {len(uniq_formats)} formatos")


if __name__ == "__main__":
    main()
